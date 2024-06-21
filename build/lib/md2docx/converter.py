import markdown2
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def convert_markdown_to_docx(markdown_text, output_file):
    # Convert markdown to HTML
    html = markdown2.markdown(markdown_text)
    
    # Create a new Document
    doc = Document()
    
    # Add the content to the document
    doc.add_heading('Converted Markdown', 0)
    
    # Convert HTML to docx
    add_html_to_doc(doc, html)
    
    # Save the document
    doc.save(output_file)

def add_html_to_doc(doc, html):
    from bs4 import BeautifulSoup
    
    soup = BeautifulSoup(html, 'html.parser')
    
    for element in soup.children:
        if element.name == 'h1':
            doc.add_heading(element.text, level=1)
        elif element.name == 'h2':
            doc.add_heading(element.text, level=2)
        elif element.name == 'h3':
            doc.add_heading(element.text, level=3)
        elif element.name == 'h4':
            doc.add_heading(element.text, level=4)
        elif element.name == 'h5':
            doc.add_heading(element.text, level=5)
        elif element.name == 'h6':
            doc.add_heading(element.text, level=6)
        elif element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Number')
        elif element.name == 'blockquote':
            para = doc.add_paragraph(element.text)
            para_format = para.paragraph_format
            para_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para_format.left_indent = Pt(36)
        elif element.name == 'code':
            pre = OxmlElement('w:pre')
            pre.text = element.text
            doc.add_paragraph(pre.text)
